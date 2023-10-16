import docx
import csv
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from tkinter import Tk, Label, Entry, Button, filedialog as fd
import tkinter as tk
from tkinter import ttk
from tkinter.messagebox import showinfo

base = Tk()
base.geometry("500x800")
base.title("Seating Plan System")

labl_0 = Label(base, text="Seating Plan System", width=20, font=("bold", 20))
labl_0.place(x=90, y=53)

labl_1 = Label(base, text="School Name", width=20, font=("bold", 10))
labl_1.place(x=80, y=130)

entry_01 = Entry(base)
entry_01.place(x=240, y=130)

labl_2 = Label(base, text="School Location", width=20, font=("bold", 10))
labl_2.place(x=68, y=180)

entry_02 = Entry(base)
entry_02.place(x=240, y=180)

labl_3 = Label(base, text="Session no.", width=20, font=("bold", 10))
labl_3.place(x=70, y=230)

entry_03 = Entry(base)
entry_03.place(x=240, y=230)

labl_4 = Label(base, text="Title of Document", width=20, font=("bold", 10))
labl_4.place(x=70, y=280)

entry_04 = Entry(base)
entry_04.place(x=240, y=280)

labl_5 = Label(base, text="Date of exam 1", width=20, font=("bold", 10))
labl_5.place(x=70, y=330)

entry_05 = Entry(base)
entry_05.place(x=240, y=330)

labl_6 = Label(base, text="Date of exam 2", width=20, font=("bold", 10))
labl_6.place(x=70, y=380)

entry_06 = Entry(base)
entry_06.place(x=240, y=380)

labl_7 = Label(base, text="Date of exam 3", width=20, font=("bold", 10))
labl_7.place(x=70, y=430)

entry_07 = Entry(base)
entry_07.place(x=240, y=430)

labl_8 = Label(base, text="Date of exam 4", width=20, font=("bold", 10))
labl_8.place(x=70, y=480)

entry_08 = Entry(base)
entry_08.place(x=240, y=480)

labl_9 = Label(base, text="Date of exam 5", width=20, font=("bold", 10))
labl_9.place(x=70, y=530)

entry_09 = Entry(base)
entry_09.place(x=240, y=530)

labl_10 = Label(base, text="Date of exam 6", width=20, font=("bold", 10))
labl_10.place(x=70, y=580)

entry_10 = Entry(base)
entry_10.place(x=240, y=580)

labl_11 = Label(base, text="Date of exam 7", width=20, font=("bold", 10))
labl_11.place(x=70, y=630)

entry_11 = Entry(base)
entry_11.place(x=240, y=630)


def filedialog():
    # create the root window
    root = tk.Tk()
    root.title("Tkinter Open File Dialog")
    root.resizable(False, False)
    root.geometry("300x150")

    def select_file():
        filetypes = (("All files", "*.csv"), ("All files", "*.*"))

        filename = fd.askopenfilename(
            title="Open the csv file", initialdir="/", filetypes=filetypes
        )

        showinfo(title="Selected File", message=filename)
        global name
        name = str(filename)

    # open button
    open_button = ttk.Button(root, text="Open a File", command=select_file)

    open_button.pack(expand=True)

    # run the application
    root.mainloop()


def get_data():
    global ltemp
    ltemp = [
        entry_01.get(),
        entry_02.get(),
        entry_03.get(),
        entry_04.get(),
        entry_05.get(),
        entry_06.get(),
        entry_07.get(),
        entry_08.get(),
        entry_09.get(),
        entry_10.get(),
        entry_11.get(),
        name,
    ]


Button(base, text="Submit", width=20, bg="brown", fg="white", command=get_data).place(
    x=170, y=750
)

filedialog()
# name =
# it will be used for displaying the registration form onto the window
base.mainloop()

f = open(ltemp[11], "r+")
data = csv.reader(f)

doc = docx.Document()

# Choosing the top most section of the page
section = doc.sections[0]

# Selecting the header
header = section.header

# Selecting the paragraph already present in
# the header section
header_para = header.paragraphs[0]
j = 1

# Adding the centred zoned header
header_para.text = f"\t{ltemp[0]}\n\t{ltemp[1]}\n\t{ltemp[2]}\n\t{ltemp[3]}"
doc.add_heading("\t\t\t\t\t\tROOM:" + str(j), 3)
table = doc.add_table(rows=1, cols=11)
table.style = "TableGrid"
table.autofit = True
# table.columns[2].width = Inches(1.5)
table.rows[0].cells[0].width = Inches(1.0)
row = table.rows[0].cells
row[0].text = "S.No "
row[1].text = "Roll NO "
row[2].text = "Name "
row[3].text = "Class & Sec "
row[4].text = ltemp[4]
row[5].text = ltemp[5]
row[6].text = ltemp[6]
row[7].text = ltemp[7]
row[8].text = ltemp[8]
row[9].text = ltemp[9]
row[10].text = ltemp[10]
i = 1

for id in data:
    # Adding a row and then adding data in it.
    row = table.add_row().cells
    table.rows[1].cells[1].width = Inches(1.0)
    # Converting id to string as table can only take string input
    row[0].text = str(id[0])
    row[1].text = str(id[1])
    row[2].text = str(id[2])
    row[3].text = str(id[3])
    if int(id[0]) == 418:
        j = 0
        i = 15

    if i == 15:
        # break;
        j = j + 1
        doc.add_page_break()
        doc.add_heading("\t\t\t\t\t\tROOM:" + str(j), 3)
        table = doc.add_table(rows=1, cols=11)
        table.style = "TableGrid"
        table.autofit = True
        # table.columns[2].width = Inches(1.5)
        table.rows[0].cells[0].width = Inches(1.0)
        row = table.rows[0].cells
        row[0].text = "S.No "
        row[1].text = "Roll Number "
        row[2].text = "Name of the Student "
        row[3].text = "Class & Sec "
        row[4].text = ltemp[4]
        row[5].text = ltemp[5]
        row[6].text = ltemp[6]
        row[7].text = ltemp[7]
        row[8].text = ltemp[8]
        row[9].text = ltemp[9]
        row[10].text = ltemp[10]
        i = 0
    i = i + 1
    for row in table.rows:
        row.height = Cm(1.2)
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(12)
    # Now save the document to a location
    section = doc.sections[0]
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.00)
    section.top_margin = Cm(1.0)

doc.save("Seating Plan.docx")
f.close()
